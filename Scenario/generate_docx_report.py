import os
import sys
import subprocess

# 1. python-docx 동적 설치 및 가져오기 시도
try:
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("python-docx 라이브러리를 설치합니다...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml, OxmlElement
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """표 셀의 배경색을 설정합니다."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """표 셀의 안쪽 여백(Padding)을 설정합니다."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC"):
    """표 테두리를 깔끔한 연한 회색으로 정돈합니다."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
            f'  <w:insideV w:val="none"/>'
            f'  <w:left w:val="none"/>'
            f'  <w:right w:val="none"/>'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def main():
    doc = docx.Document()
    
    # 2. 문서 마진 설정 (좁게 설정하여 보고서 가독성 확대)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # 3. 색상 팔레트 설정 (클래식 블랙 & 화이트 테마)
    COLOR_PRIMARY = RGBColor(17, 24, 39)      # 다크 그레이/블랙
    COLOR_SECONDARY = RGBColor(55, 65, 81)     # 미디움 그레이
    COLOR_TEXT = RGBColor(17, 24, 39)          # 본문 텍스트
    COLOR_MUTED = RGBColor(107, 114, 128)      # 연한 그레이 (메타 데이터)
    
    # 4. 기본 폰트 설정
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    font.color.rgb = COLOR_TEXT
    
    # 5. 타이틀 페이지 구성
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("로컬 AI 기반 시나리오 어시스턴트: 최종 결과 보고서\n")
    title_run.font.name = '맑은 고딕'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    sub_run = title_p.add_run("Local LLM Fine-Tuning & Knowledge Graph RAG System for Scenario Writing")
    sub_run.font.name = '맑은 고딕'
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_MUTED
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    meta_run = meta_p.add_run("작성일자: 2026년 06월 18일\n프로젝트명: 시나리오 어시스턴트 (Scenario Assistant)\n발표자: C077036 차현민")
    meta_run.font.name = '맑은 고딕'
    meta_run.font.size = Pt(9.5)
    meta_run.font.color.rgb = COLOR_MUTED
    
    doc.add_paragraph("-" * 80)
    
    # 공통 헤더 추가 함수
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_bullet_item(bold_prefix, text_content):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run_bold = p.add_run(bold_prefix)
        run_bold.font.name = '맑은 고딕'
        run_bold.font.bold = True
        run_bold.font.color.rgb = COLOR_TEXT
        
        run_text = p.add_run(text_content)
        run_text.font.name = '맑은 고딕'
        run_text.font.color.rgb = COLOR_TEXT
        return p

    # 1. 연구 배경 및 문제 정의
    add_section_header("1. 연구 배경 및 문제 정의 (Background & Problem Statement)")
    
    p = doc.add_paragraph()
    p.add_run(
        "과거 소설이나 게임 시나리오 라이팅에서 세계관은 단순히 주인공들이 움직이는 뒷배경(Setting)에 불과했습니다. "
        "그러나 현대 서사 콘텐츠(웹소설, 멀티엔딩 RPG, 웹툰 등)는 방대한 회차와 정교한 타임라인을 갖추게 되면서, "
        "세계관의 일관성과 디테일이 독자/유저의 몰입감을 결정하는 핵심 요소로 자리 잡았습니다.\n\n"
        "세계관의 규모가 확장됨에 따라 작가는 수십 명의 등장인물 정보, 수많은 배경 장소, 복잡한 인과관계의 플롯을 기억하고 통제해야 합니다. "
        "이 과정에서 필연적으로 다음과 같은 한계가 발생합니다."
    )
    
    add_bullet_item("• 휴먼 에러 (Human Error): ", "연재 회차가 길어질수록 과거 설정을 망각하고 상충되는 인물 성향이나 모순된 타임라인을 서술하는 설정 붕괴 현상이 잦아짐.")
    add_bullet_item("• 관리 파편화: ", "인물 메모장, 장소 텍스트 파일, 관계도 그림 파일 등 여러 프로그램에 세계관 정보가 파편화되어 집필 중에 수시로 탐색해야 하므로 집필 흐름이 단절됨.")

    # 2. 해결 방안 및 핵심 기능
    add_section_header("2. 해결 방안 및 핵심 기능 (Solutions & Core Features)")
    
    p = doc.add_paragraph()
    p.add_run(
        "창작자가 모든 방대한 설정을 머릿속에 기억하고 유지하는 것은 불가능하다는 한계를 인정해야 합니다. "
        "창작자는 오직 \"원문 집필(텍스트 창작)\"에만 몰입하고, 설정 추출, 병합, 무결성 검증과 같은 부차적인 관리는 AI가 대신하여 자동화한다면 생산성과 결과물의 품질이 극대화될 것입니다. "
        "AI를 창작의 주체가 아닌, 창작을 보조하는 최적의 비서(Assistant)로 설계합니다.\n\n"
        "시스템 내 핵심 구현 기능은 다음과 같습니다."
    )
    
    add_bullet_item("1) AI 스토리 데이터 구조화 & 개체 태깅: ", "작가가 집필한 원문을 복사·붙여넣기 하면, 로컬 모델이 원문 속 등장인물(Characters), 장소(Locations), 사건(Plots)을 분류하고, 각 개체 간의 인과관계를 관계망(Graph) 형태로 정밀 추출합니다.")
    add_bullet_item("2) 세계관 설정 무결성 검증 (Integrity Agent Check): ", "등록된 모든 세계관 정보를 텍스트로 요약·직렬화하여 AI 검증 에이전트에 통째로 주입합니다. 에이전트는 캐릭터의 사망 여부, 장소 이동의 물리학적 한계, 시간대 상충 등을 크로스 체크하여 UI에 경고 카드로 시각화합니다.")
    add_bullet_item("3) 자동 타임라인 시각화 및 양방향 관계 관리 대시보드: ", "추출된 사건들을 년도별 시간선 순서로 순차 시각화하고, 인물·장소·플롯을 카드형 대시보드로 관리하며 상세 속성 창에서 관계 설정 시 자동으로 실시간 역방향 관계를 매핑하는 유기적 에디터 환경을 제공합니다.")

    # 3. 구현 전략
    add_section_header("3. 구현 전략 (Implementation Strategy: Local AI + Web App)")
    
    p = doc.add_paragraph()
    p.add_run(
        "• 아이디어 보안 노출 (Security): 상용 클라우드 대형 언어 모델 API를 사용할 경우, 미출간된 고유 설정과 작품 원문이 외부 서버로 전송됩니다. "
        "이는 개인 작가뿐만 아니라 상업적 지식재산권(IP)을 가진 제작사 관점에서도 결코 허용될 수 없는 치명적인 보안 취약점입니다. 따라서 100% 오프라인 구동 구조를 채택했습니다.\n"
        "• 로컬 서빙 경량화 (FastAPI + Ollama): 기존의 PyTorch 환경을 쌩으로 로드하여 실행하는 것은 개인 기기에서 지나치게 무겁습니다. "
        "본 프로젝트는 Ollama 서비스와 Python FastAPI를 백엔드로 조합하여, 사용자의 개인 컴퓨터에서도 지연 없이 동작하는 가벼운 개인 비서 컨셉의 인프라를 구축하였습니다."
    )

    # 4. 기존 로컬 모델의 문제점 및 SFT
    add_section_header("4. 기존 로컬 모델의 문제점 및 SFT(QLoRA) 도입")
    
    p = doc.add_paragraph()
    p.add_run(
        "기존 로컬 모델은 크게 (1) JSON 포맷 구조를 완전히 준수하지 못해 파싱 오류를 내는 형식 파괴 문제, "
        "(2) 소설 고유의 문맥적 개체 요소를 정확하게 파악해내지 못하는 정보 추출력 부실 문제가 식별되었습니다.\n\n"
        "이를 해결하고자 지식 증류(Knowledge Distillation) 전략을 도입하여 성능이 우수한 Claude 3.5 Haiku 모델로 "
        "1930년대 한국 소설 데이터셋(200 샘플)의 세계관 구조 정답지를 먼저 추출(Pseudo-Gold Standard)하고, "
        "경량 로컬 모델들이 이를 그대로 습득하여 추출율과 스키마 준수율을 동시에 개선하도록 SFT QLoRA 학습을 진행했습니다."
    )

    # 5. 파인튜닝 학습 파라미터 표
    add_section_header("5. 파인튜닝 학습 파라미터 (Hyperparameters)")
    
    p = doc.add_paragraph("경량 로컬 온디바이스 모델의 효율적 학습을 위해 구성한 최적의 하이퍼파라미터 세팅입니다.")
    
    params = [
        ("Fine-Tuning 기법", "QLoRA (4-bit)", "VRAM OOM 방지 및 양자화 계산 정밀도(NF4) 활용"),
        ("Base Precision", "bf16 (Bfloat16)", "Ampere 계열 이상 GPU의 연산 정밀도 최적화"),
        ("LoRA Rank (r)", "32", "단순 출력 형태 모사를 넘어 복잡한 데이터 인과관계 습득"),
        ("LoRA Alpha (alpha)", "64", "Rank 크기 대비 안정적인 스케일링 팩터 적용"),
        ("Target Modules", "all-linear", "Attention 뿐만 아니라 MLP 레이어 전체 튜닝으로 지식 암기"),
        ("Epochs", "4", "200 샘플 규모에서 과적합을 방지하고 완벽한 포맷 고정"),
        ("Batch Size", "1", "VRAM 한계를 감안한 Micro Batch Size 설정"),
        ("Gradient Accumulation", "8", "실질 학습 배치 사이즈를 8(1 x 8)로 유지하여 안정성 확보"),
        ("Learning Rate", "2e-4", "Cosine Scheduler와 연동된 안정적인 수렴 속도"),
        ("LR Scheduler", "cosine", "학습 후반부의 세밀한 가중치 미세 조율"),
        ("Optimizer", "paged_adamw_8bit", "메모리 페이지 오프로딩 기법을 통해 Peak VRAM 제어"),
        ("Sequence Length", "1024", "소설 원문 및 프롬프트 주입 맥락 길이 확보")
    ]
    
    table_p = doc.add_table(rows=1, cols=3)
    table_p.style = 'Table Grid'
    set_table_borders(table_p)
    
    hdr_cells = table_p.rows[0].cells
    hdr_cells[0].text = '하이버파라미터'
    hdr_cells[1].text = '설정 값'
    hdr_cells[2].text = '비고 / 최적화 사유'
    for cell in hdr_cells:
        set_cell_background(cell, "F3F4F6") # 연한 그레이 배경
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = COLOR_TEXT # 흰색 대신 검은색 텍스트
        cell.paragraphs[0].runs[0].font.name = '맑은 고딕'
        
    for p_name, p_val, p_desc in params:
        row = table_p.add_row()
        row.cells[0].text = p_name
        row.cells[1].text = p_val
        row.cells[2].text = p_desc
        for cell in row.cells:
            set_cell_background(cell, "FFFFFF") # 화이트 배경
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.name = '맑은 고딕'
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph() # spacing

    # 6. 실험 결과 및 성능 지표 비교
    add_section_header("6. 실험 결과 및 성능 지표 비교 (Evaluation Results)")
    
    p = doc.add_paragraph("SFT 진행 전(Base)과 진행 후(SFT)의 지표를 200개의 평가 청크 데이터로 정밀 측정한 비교 지표 결과입니다.")
    
    eval_results = [
        ("Gemma Base (2.6B)", "53.00%", "665", "2.04", "0.1739"),
        ("Gemma SFT Opt (4B)", "97.95%", "1,564", "3.76", "0.4972"),
        ("DeepSeek SFT (1.5B)", "100.00%", "2,065", "3.51", "*0.0054"),
        ("Qwen SFT (4B)", "100.00%", "1,415", "3.87", "0.5294"),
        ("EXAONE SFT (1.2B)", "100.00%", "1,508", "3.72", "0.4951")
    ]
    
    table_e = doc.add_table(rows=1, cols=5)
    table_e.style = 'Table Grid'
    set_table_borders(table_e)
    
    hdr_e = table_e.rows[0].cells
    headers = ['모델명', '스키마 준수율', '총 추출 개체 수', '개체당 평균 묘사', 'F1-Score']
    for idx, name in enumerate(headers):
        hdr_e[idx].text = name
        set_cell_background(hdr_e[idx], "F3F4F6") # 연한 그레이 배경
        set_cell_margins(hdr_e[idx], top=120, bottom=120, left=150, right=150)
        hdr_e[idx].paragraphs[0].runs[0].font.bold = True
        hdr_e[idx].paragraphs[0].runs[0].font.color.rgb = COLOR_TEXT # 검은색 텍스트
        hdr_e[idx].paragraphs[0].runs[0].font.name = '맑은 고딕'
        
    for m_name, s_rate, t_ent, a_det, f_score in eval_results:
        row = table_e.add_row()
        row.cells[0].text = m_name
        row.cells[1].text = s_rate
        row.cells[2].text = t_ent
        row.cells[3].text = a_det
        row.cells[4].text = f_score
        for idx, cell in enumerate(row.cells):
            set_cell_background(cell, "FFFFFF" if idx != 4 or f_score != "0.5294" else "F9FAFB") # 강조 행에 미세한 회색 처리
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.paragraphs[0].runs[0].font.name = '맑은 고딕'
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)
            if idx == 4 and f_score == "0.5294":
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = COLOR_PRIMARY

    doc.add_paragraph()

    # 콜아웃 경고박스 추가 (Single-cell table 이용)
    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    callout_cell = callout.rows[0].cells[0]
    set_cell_background(callout_cell, "F9FAFB") # 라이트 그레이 배경
    set_cell_margins(callout_cell, top=140, bottom=140, left=200, right=200)
    # 테두리 차분한 다크 그레이로 설정
    tcPr = callout_cell._element.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="0" w:color="4B5563"/><w:top w:val="none"/><w:right w:val="none"/><w:bottom w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)
    
    cp = callout_cell.paragraphs[0]
    c_run = cp.add_run("[!] DeepSeek SFT의 F1-Score 왜곡 지표 설명 (*0.0054):\n")
    c_run.font.name = '맑은 고딕'
    c_run.font.bold = True
    c_run.font.color.rgb = COLOR_PRIMARY
    c_run.font.size = Pt(9.5)
    
    c_run2 = cp.add_run(
        "DeepSeek-R1-Distill-Qwen-1.5B 모델은 추론(Reasoning) 특화 모델로서 JSON 스키마 일관성은 100% 보장하며 엔티티도 대량 추출했으나, "
        "추론 과정의 출력 필터링 오류로 인한 결측치(Empty values)가 생성되거나, 원본 한국어 개체명 대조 시 자모 결합 오차 및 형태소 편차로 인해 "
        "단순 문자열 매칭 기반의 F1-score 측정식에서 과소평가된 한계가 존재합니다."
    )
    c_run2.font.name = '맑은 고딕'
    c_run2.font.size = Pt(9)
    c_run2.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()

    # 학습 모델별 특징 리스트
    add_bullet_item("1) Qwen SFT (4B): ", "가장 완성도 높은 JSON 구조화 출력 안정성과 정답 재현력을 보여주었으며, F1-Score 0.5294로 로컬 모델 중 최고점을 달성했습니다.")
    add_bullet_item("2) Gemma SFT Opt (4B): ", "한국어 문장의 함축적 지식 및 인물 행동 묘사 추출도가 매우 탁월했으나, 미세한 JSON 태그 에러율(약 2.05%)이 식별되어 정규식 보정 필터가 추가되었습니다.")
    add_bullet_item("3) EXAONE SFT (1.2B): ", "LGAI 오픈소스 한국어 특화 모델로서 1B급의 초경량 체급임에도 불구하고 스키마 준수율 100%와 우수한 묘사력(3.72)을 기록해 리소스가 극도로 제한된 온디바이스 최적의 후보임을 증명했습니다.")

    # 7. 프로젝트의 한계 및 결론
    add_section_header("7. 프로젝트의 한계 및 결론 (Limitations & Conclusion)")
    
    p = doc.add_paragraph()
    p.add_run(
        "• 데이터셋 크기의 절대적 부족과 성능 한계: 본 프로젝트를 통해 로컬 모델 파인튜닝 시 스키마 준수율을 100%로 끌어올리고 추출 지표를 베이스 모델 대비 3배 이상 대폭 상승시키는 등 소기의 성과를 달성했습니다. "
        "그러나 절대적인 F1-Score 성능 수치(최대 0.52급)는 여전히 상용 거대 모델(Claude 3.5)의 완성도를 완벽히 대체하기에는 아쉬운 격차가 존재합니다. 이 원인을 분석한 결과, 200개 샘플이라는 정답 학습 데이터셋(Dataset Size)의 절대적 양 부족이 가장 큰 요인으로 판단됩니다.\n"
        "• 데이터셋 개선을 가로막는 자본의 현실적 장벽: 현실적으로 성능을 상용 수준으로 끌어올리려면 수천 장에서 수만 장에 이르는 장편 소설 데이터를 추가로 구축하여 더 큰 파인튜닝 비용과 GPU 연산 자원을 투입해야 합니다. "
        "그러나 개인 창작자 혹은 소규모 인력으로서 이 이상의 막대한 구축 자금을 상용 API 추출 및 학습 장비 대여에 투입하는 것은 현실적으로 불가능한 장벽이었습니다.\n"
        "• [개인적 통찰] 자본의 전쟁터가 된 AI 시장 속 개인의 역할: 현재 글로벌 AI 시장은 천문학적인 자본력과 초대형 데이터 센터를 지닌 빅테크 기업들에 의해 주도되는 거대한 자본 전쟁터입니다. "
        "이 거대한 흐름 속에서 작은 개인이나 소규모 개발자가 과연 무엇을 기여하고 변화시킬 수 있을지에 대한 회의감과 의문이 깊이 남는 것은 사실입니다.\n\n"
        "그럼에도 불구하고, 보안성을 철저히 지키며 사용자의 곁에서 묵묵히 보조하는 \"온디바이스 로컬 창작 비서\"의 가능성을 확인한 것만으로도 의미 있는 도전이었습니다. "
        "거대 플랫폼이 모든 데이터를 독점하려는 구조 속에서, 로컬 기기 내에서 스스로 작동하는 고유한 특화 AI는 자본의 전쟁 틈새 속에서 개인이 나아갈 수 있는 또 하나의 작은 돌파구가 될 수 있을 것입니다."
    )

    # 8. 파일 저장
    output_path = "Scenario_Assistant_Results_Report.docx"
    doc.save(output_path)
    print(f"성공적으로 Word 보고서 생성 완료: {os.path.abspath(output_path)}")

if __name__ == "__main__":
    main()
